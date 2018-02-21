import PyPDF2
import os

os.getcwd()
os.chdir('week_1')
os.listdir()
os.rename('Psychobiology_literature_week1.pdf', '2.pdf')

# work with pdfs
pdf = open('1.pdf', 'rb')
pdfreader = PyPDF2.PdfFileReader(pdf)
#page numbers
pdfreader.numPages
page = pdfreader.getPage(1)
page.extractText()
pdf.close()

#I stop here because I don't see any use cases for me at this moment.

from docx import Document
os.getcwd()
os.chdir('proposal')
os.listdir()
doc = Document('exp.docx')
len(doc.paragraphs)
doc.paragraphs[5].text
len(doc.paragraphs[6].runs)
doc.paragraphs[6].runs[1].text

#a function to extract only the text 
def getText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

os.getcwd()
print(getText('exp.docx'))


def rb_hang_ind(filename):
    '''
    Takes a docx filename out of cwd as string, searches for a paragraph that equals 'References' in that document and adds hanging ident and double spacing to all following paragraphs.
    '''
    #import modules
    from docx import Document
    from docx.shared import Pt, Inches, Length
    #load document 
    doc = Document(filename)
    #number of paragraphs
    n = len(doc.paragraphs)
    #find the reference list 
    for p in range(n):
        if doc.paragraphs[p].text == '2d. Literature references ':
            start = p+1 
        if doc.paragraphs[p].text == '2e. Time Plan ':
            end = p
    #add hanging ident to all following lines (dont use if reflist is not last part)
    for p in range(start, end):
        doc.paragraphs[p].paragraph_format.left_indent = Pt(36)
        doc.paragraphs[p].paragraph_format.first_line_indent = Inches(-0.5)
        doc.paragraphs[p].paragraph_format.line_spacing = 2.0

    doc.save('hang_ind_{}'.format(filename))

rb_hang_ind('first_version.docx')

import os
os.getcwd()
os.chdir('block_2/developmental_psychopathology/term_paper')
os.listdir()
rb_hang_ind('final.docx')

for i in range(1,9):
    if i < 6:
        os.makedirs('week_{}/hw{}'.format(i,i))
    else:
        os.mkdir('week_{}'.format(i))
        